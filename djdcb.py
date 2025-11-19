# djdcb_final.py
"""
最终版：按你的严格要求实现 DXF(TK 层) -> Word 填表
规则：
 - 只从 TK 图层提取 J 点和两位小数边长（1.00~100.00）
 - 点按编号升序排序（J1,J2,...）
 - 每对相邻点 (Ji, J{i+1}) 的边长分配给 Ji（循环闭合：Jn->J1）
 - 配对通过对段中点找最近的未用边长，若多个候选优先位于“右侧”的文本
 - 写 Word：优先写入已有 J 行（只写边长列）-> 填表中空位 -> 不够则分页复制表格 -> 最后追加行
 - 写单元格时尽量保留原样式：替换第一个 run 的 text，不重建 cell
"""

import re
import math
import copy
import os
from collections import namedtuple

import ezdxf
from docx import Document
from docx.enum.text import WD_BREAK

# ------------------ 配置：修改为你本地路径 ------------------
DXF_PATH = r"D:\cadtu\D20XTT20250055_宗地图.dxf"            # 你的 DXF 文件路径
TEMPLATE_DOCX = r"D:\cadtu\（样版）新版地籍调查表.docx"     # 模板路径
OUTPUT_DOCX = r"D:\cadtu\填充后的新版地籍调查表_final.docx"  # 输出路径
LAYER_NAME = "TK"
# -----------------------------------------------------------

PointText = namedtuple("PointText", ["text", "x", "y"])

# regex
_RE_J = re.compile(r'[Jj]\s*0*([0-9]{1,4})')       # find J number anywhere
_RE_DEC = re.compile(r'(\d{1,3}\.\d{2})')         # two decimals like 2.10

def normalize_j(s: str) -> str:
    if not s:
        return ""
    m = _RE_J.search(s)
    if m:
        return f"J{int(m.group(1))}"
    return ""

def extract_decimal(s: str):
    if not s:
        return None
    m = _RE_DEC.search(s.replace("，", ",").replace("。", "."))
    if m:
        return m.group(1)
    return None

def safe_coords(e):
    """safe extraction of insert coordinates, fallback to 0,0"""
    try:
        ins = e.dxf.insert
        return float(ins[0]), float(ins[1])
    except Exception:
        try:
            ins = e.insert
            return float(ins[0]), float(ins[1])
        except Exception:
            return 0.0, 0.0

def extract_from_dxf(dxf_path, layer_name="TK"):
    """Extract J points and decimal lengths (only from layer_name)"""
    try:
        doc = ezdxf.readfile(dxf_path)
    except Exception as e:
        raise RuntimeError(f"读取 DXF 失败: {e}")

    msp = doc.modelspace()
    raw = []
    for e in msp:
        try:
            t = e.dxftype()
        except Exception:
            continue
        if t == "TEXT":
            txt = getattr(e.dxf, "text", "") or ""
            x, y = safe_coords(e)
            layer = getattr(e.dxf, "layer", "")
            raw.append({"text": txt, "x": x, "y": y, "layer": layer})
        elif t == "MTEXT":
            txt = getattr(e, "text", "") or getattr(e.dxf, "text", "") or ""
            x, y = safe_coords(e)
            layer = getattr(e.dxf, "layer", "")
            raw.append({"text": txt, "x": x, "y": y, "layer": layer})
        elif t == "INSERT":
            # iterate attributes if present
            try:
                base_x, base_y = safe_coords(e)
            except:
                base_x, base_y = 0.0, 0.0
            layer = getattr(e.dxf, "layer", "")
            try:
                for att in e.attribs:
                    try:
                        a_txt = att.dxf.text
                    except Exception:
                        a_txt = getattr(att, "text", "") or ""
                    try:
                        a_x, a_y = safe_coords(att)
                    except:
                        a_x, a_y = base_x, base_y
                    raw.append({"text": a_txt, "x": a_x, "y": a_y, "layer": layer})
            except Exception:
                continue
        else:
            continue

    # Filter TK layer for both J points and lengths (per your description)
    points = []
    lengths = []
    for it in raw:
        txt = str(it["text"]).strip()
        x, y = float(it["x"]), float(it["y"])
        layer = it.get("layer", "")
        # only consider J if text contains J and entity is on TK as you specified
        # The user explicitly said "提取TK图层的带有J的和2.10两位小数的数"
        # So require layer == LAYER_NAME for both categories
        if layer == layer_name:
            j = normalize_j(txt)
            if j:
                points.append(PointText(j, x, y))
                continue
            dec = extract_decimal(txt)
            if dec:
                try:
                    fv = float(dec)
                    if 1.00 <= fv <= 100.00:
                        lengths.append(PointText(dec, x, y))
                except:
                    pass
    # deduplicate points by label, keep first occurrence
    seen = {}
    pts_unique = []
    for p in points:
        if p.text not in seen:
            seen[p.text] = True
            pts_unique.append(p)
    # sort by index numeric ascending
    try:
        pts_unique.sort(key=lambda p: int(p.text[1:]))
    except:
        pass

    return pts_unique, lengths, raw

# geometric helpers
def midpoint(a: PointText, b: PointText):
    return ((a.x + b.x)/2.0, (a.y + b.y)/2.0)

def vec(a, b):
    return (b.x - a.x, b.y - a.y)

def right_vector(dx, dy):
    # rotate (dx,dy) clockwise by 90 deg to get right side vector
    return (dy, -dx)

def dot(u, v):
    return u[0]*v[0] + u[1]*v[1]

def dist_sq(x1, y1, x2, y2):
    dx = x1 - x2; dy = y1 - y2
    return dx*dx + dy*dy

def assign_lengths_by_adjacent(points, lengths):
    """
    For each adjacent pair (Ji, J{i+1}), including (Jn,J1),
    compute midpoint and pick nearest unused length point.
    Prefer candidate that is located on the segment's right side (dot > 0).
    Return mapping { 'Ji': 'len' } for i=1..n (last maps Jn -> length between Jn and J1).
    """
    mapping = {}
    if not points or not lengths:
        return mapping

    # Build list of candidate lengths with coords and used flag
    cand = [{"val": l.text, "x": l.x, "y": l.y, "used": False} for l in lengths]

    n = len(points)
    for i in range(n):
        a = points[i]
        b = points[(i+1) % n]  # next, wrap-around
        mx, my = midpoint(a, b)
        # compute segment vector and right vector
        dx, dy = vec(a, b)
        rx, ry = right_vector(dx, dy)
        # find nearest candidate(s) among unused
        best_idx = None
        best_ds = None
        best_rightness = None
        for idx, c in enumerate(cand):
            if c["used"]:
                continue
            ds = dist_sq(mx, my, c["x"], c["y"])
            # compute rightness: dot product of (c - midpoint) with right vector
            vx = c["x"] - mx; vy = c["y"] - my
            rightness = dot((vx, vy), (rx, ry))
            # choose winner: prefer smaller ds; if ds similar, prefer positive rightness
            if best_idx is None:
                best_idx, best_ds, best_rightness = idx, ds, rightness
            else:
                # if new candidate significantly closer, take it
                if ds < best_ds - 1e-6:
                    best_idx, best_ds, best_rightness = idx, ds, rightness
                elif abs(ds - best_ds) <= 1e-6:
                    # tie: prefer positive rightness
                    if (rightness > 0 and best_rightness <= 0):
                        best_idx, best_ds, best_rightness = idx, ds, rightness
        if best_idx is not None:
            mapping[a.text] = cand[best_idx]["val"]
            cand[best_idx]["used"] = True
        else:
            # no candidate found: leave blank
            mapping[a.text] = ""
    return mapping

# Word helpers
def set_cell_text_preserve(cell, text):
    """Replace only first run text, clear other runs to avoid residual text.
       If cell has multiple paragraphs/runs this will attempt to preserve paragraph style/formatting."""
    text = "" if text is None else str(text)
    try:
        paras = cell.paragraphs
    except Exception:
        cell.text = text
        return
    if not paras:
        p = cell.add_paragraph()
        p.add_run(text)
        return
    p0 = paras[0]
    if p0.runs:
        p0.runs[0].text = text
        for r in p0.runs[1:]:
            r.text = ""
    else:
        p0.add_run(text)
    # clear other paragraphs contents to avoid leftovers (but keep paragraph objects)
    for p in paras[1:]:
        for r in p.runs:
            r.text = ""

def find_table_and_columns(doc):
    """Find the table containing headers for point and length; robustly determine columns."""
    for tbl in doc.tables:
        max_check = min(10, len(tbl.rows))
        header_row = None
        col_point = None
        col_length = None
        for r in range(max_check):
            row = tbl.rows[r]
            for c, cell in enumerate(row.cells):
                txt_flat = cell.text.replace("\n", "").replace(" ", "")
                lt = cell.text.lower()
                if any(k in txt_flat for k in ("界址点号", "界址编号", "点号", "界址")) and col_point is None:
                    col_point = c
                    header_row = r
                if ("边长" in cell.text or "长度" in cell.text or "边 长" in cell.text) and col_length is None:
                    col_length = c
                    header_row = r if header_row is None else header_row
            if header_row is not None and (col_point is not None or col_length is not None):
                break
        if header_row is not None:
            # if col_length not found, try heuristic: col_point + 1 or search header row for '边'
            num_cols = len(tbl.rows[header_row].cells)
            if col_point is None:
                col_point = 0
            if col_length is None:
                for c, cell in enumerate(tbl.rows[header_row].cells):
                    if "边" in cell.text or "长" in cell.text:
                        col_length = c
                        break
            if col_length is None:
                if (col_point + 1) < num_cols:
                    col_length = col_point + 1
                else:
                    col_length = num_cols - 1
            # sanitize
            col_point = max(0, min(col_point, num_cols - 1))
            col_length = max(0, min(col_length, num_cols - 1))
            return tbl, header_row, col_point, col_length
    return None, None, None, None

def clone_table_to_new_page(doc, table, preserve_header_rows):
    """Insert page break and deepcopy table XML to replicate the same table on new page; clear data area text."""
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    new_tbl_xml = copy.deepcopy(table._tbl)
    doc._body._element.append(new_tbl_xml)
    new_tbl = doc.tables[-1]
    # clear data rows (preserve header rows)
    for ri, row in enumerate(new_tbl.rows):
        if ri < preserve_header_rows:
            continue
        for cell in row.cells:
            set_cell_text_preserve(cell, "")
    return new_tbl

def fill_word_table(doc, table, header_row, col_point, col_length, points, mapping_len):
    data_start = header_row + 1
    # validation of indexes
    num_cols = len(table.rows[header_row].cells)
    if not (0 <= col_point < num_cols):
        col_point = 0
    if not (0 <= col_length < num_cols):
        col_length = min(col_point + 1, num_cols - 1)

    # 1) first write lengths into existing rows that already have J label (only write length column)
    written = set()
    for r in range(data_start, len(table.rows)):
        try:
            existing = table.cell(r, col_point).text
        except Exception:
            existing = ""
        norm = normalize_j(existing)
        if norm and norm in mapping_len:
            set_cell_text_preserve(table.cell(r, col_length), mapping_len[norm])
            written.add(norm)

    # 2) fill remaining J labels into empty point-column rows (in ascending J order)
    j_order = [p.text for p in points]
    remaining = [j for j in j_order if j not in written]
    if remaining:
        empty_rows = []
        for r in range(data_start, len(table.rows)):
            try:
                t = table.cell(r, col_point).text.strip()
            except Exception:
                t = ""
            if not t:
                empty_rows.append(r)
        for i, j in enumerate(remaining):
            if i < len(empty_rows):
                r = empty_rows[i]
                set_cell_text_preserve(table.cell(r, col_point), j)
                set_cell_text_preserve(table.cell(r, col_length), mapping_len.get(j, ""))
                written.add(j)
            else:
                break
        remaining = [j for j in remaining if j not in written]

    # 3) if still remaining, clone table to new page(s) and fill there
    preserve_header_rows = header_row + 1
    current_table = table
    while remaining:
        current_table = clone_table_to_new_page(doc, current_table, preserve_header_rows)
        # find empties in new table
        empty_rows_new = []
        for r in range(preserve_header_rows, len(current_table.rows)):
            try:
                t = current_table.cell(r, col_point).text.strip()
            except Exception:
                t = ""
            if not t:
                empty_rows_new.append(r)
        filled_here = 0
        for i, j in enumerate(remaining):
            if i < len(empty_rows_new):
                r = empty_rows_new[i]
                set_cell_text_preserve(current_table.cell(r, col_point), j)
                set_cell_text_preserve(current_table.cell(r, col_length), mapping_len.get(j, ""))
                written.add(j)
                filled_here += 1
            else:
                break
        remaining = [j for j in remaining if j not in written]
        if filled_here == 0:
            # no empty rows in cloned table: break to append rows
            break

    # 4) if still remaining, append rows to the last table
    if remaining:
        last_table = current_table
        for j in remaining:
            new_row = last_table.add_row()
            if col_point < len(new_row.cells):
                set_cell_text_preserve(new_row.cells[col_point], j)
            if col_length < len(new_row.cells):
                set_cell_text_preserve(new_row.cells[col_length], mapping_len.get(j, ""))
            written.add(j)
        remaining = []

    return written

# --- main process ---
def process(dxf_path, template_docx, output_docx, layer_name="TK"):
    print("STEP 1: 从 DXF 提取 (只读取 TK 图层中的 J 与 两位小数)...")
    points, lengths, raw = extract_from_dxf(dxf_path, layer_name=layer_name)
    print(f"  提取到点号 count={len(points)} 示例前40: {[p.text for p in points[:40]]}")
    print(f"  提取到边长 count={len(lengths)} 示例前40: {[l.text for l in lengths[:40]]}")

    print("\nSTEP 2: 按相邻点配对并分配边长（每对段中点选最近的边长，优先右侧）...")
    mapping_len = assign_lengths_by_adjacent(points, lengths)
    mapped_count = sum(1 for v in mapping_len.values() if v)
    print(f"  分配到的边长数: {mapped_count} / {len(points)} （示例前40）:")
    i = 0
    for k in sorted(mapping_len.keys(), key=lambda s: int(s[1:]) if s.startswith("J") else 999999):
        if i >= 40:
            break
        print(f"    {k} -> {mapping_len[k]}")
        i += 1

    print("\nSTEP 3: 打开 Word 模板并定位界址标示表 ...")
    doc = Document(template_docx)
    table, header_row, col_point, col_length = find_table_and_columns(doc)
    if table is None:
        raise RuntimeError("未找到包含 '界址点号' 或 '边长' 的表格，请检查模板表头。")
    print(f"  找到表格: header_row={header_row}, col_point={col_point}, col_length={col_length}, cols_in_header={len(table.rows[header_row].cells)}")

    print("\nSTEP 4: 填写 Word 表格（严格保留表头与样式） ...")
    written = fill_word_table(doc, table, header_row, col_point, col_length, points, mapping_len)
    print(f"  已写入点数: {len(written)} (示例前30: {list(written)[:30]})")

    doc.save(output_docx)
    print("\nSTEP 5: 保存输出 ->", os.path.abspath(output_docx))
    return {
        "points_extracted": [p.text for p in points],
        "lengths_extracted": [l.text for l in lengths],
        "mapping_len": mapping_len,
        "written": written,
        "output": os.path.abspath(output_docx)
    }

if __name__ == "__main__":
    res = process(DXF_PATH, TEMPLATE_DOCX, OUTPUT_DOCX, layer_name=LAYER_NAME)
    print("\n=== summary ===")
    print("points:", len(res["points_extracted"]))
    print("lengths:", len(res["lengths_extracted"]))
    print("mapped:", sum(1 for v in res["mapping_len"].values() if v))
    print("written:", len(res["written"]))
