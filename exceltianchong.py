import pandas as pd
from docx import Document
import os
import re


def replace_text_in_doc(doc, old_text, new_text):
    """替换文档中所有出现的指定文本，保持原格式"""
    # 处理段落中的文本
    for para in doc.paragraphs:
        if old_text in para.text:
            # 逐段替换以保留格式
            for run in para.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, str(new_text))

    # 处理表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, str(new_text))


def generate_documents(excel_path, template_path, output_dir='不动产登记文件'):
    """
    从Excel读取数据并填入Word模板
    - 忽略Excel第二行（描述行）
    - 以数据001字段值作为生成文件的名称
    - 仅处理前5条有效数据
    """
    # 验证文件路径是否存在
    if not os.path.exists(excel_path):
        print(f"错误：Excel文件路径不存在 - {excel_path}")
        return
    if not os.path.exists(template_path):
        print(f"错误：模板文件路径不存在 - {template_path}")
        return

    # 创建输出目录
    os.makedirs(output_dir, exist_ok=True)
    print(f"输出目录：{os.path.abspath(output_dir)}")

    try:
        # 读取Excel数据：跳过第二行（索引1），取前5条数据
        df = pd.read_excel(
            excel_path,
            sheet_name='Sheet1',
            header=0,  # 第一行为表头
            skiprows=[1],  # 跳过第二行描述信息
            nrows=5  # 仅读取5条测试数据
        )
        print(f"成功读取Excel数据，共 {len(df)} 条记录")
        print("Excel列名：", df.columns.tolist())

    except Exception as e:
        print(f"读取Excel失败：{str(e)}")
        return

    # 字段映射（Excel列名与Word模板占位符完全一致）
    field_mapping = {
        '数据002': '数据002',
        '数据003': '数据003',
        '数据004': '数据004',
        '数据005': '数据005',
        '数据006': '数据006',
        '数据001': '数据001'  # 包含数据001用于填充和命名
    }

    # 用于命名文件的列名（数据001）
    name_column = '数据001'

    # 检查必要列是否存在
    missing_columns = [col for col in field_mapping.keys() if col not in df.columns]
    if missing_columns:
        print(f"警告：Excel中缺少以下必要列，可能导致填充失败：{missing_columns}")

    # 处理每条数据
    for index, row in df.iterrows():
        try:
            # 打开模板
            doc = Document(template_path)

            # 替换所有字段
            for excel_field, doc_placeholder in field_mapping.items():
                if excel_field not in df.columns:
                    continue  # 跳过不存在的列
                # 获取单元格值，空值处理为空白
                cell_value = row[excel_field] if pd.notna(row[excel_field]) else ''
                replace_text_in_doc(doc, doc_placeholder, cell_value)

            # 生成文件名（使用数据001的值）
            if name_column in df.columns and pd.notna(row[name_column]):
                raw_name = str(row[name_column])
                # 移除文件名中的非法字符
                valid_name = re.sub(r'[\\/*?:"<>|]', '', raw_name)
                filename = f"{valid_name}.docx"
            else:
                #  fallback：使用索引作为文件名
                filename = f"文档_{index + 1}.docx"

            # 保存文件
            output_path = os.path.join(output_dir, filename)
            doc.save(output_path)
            print(f"[{index + 1}/{len(df)}] 生成成功：{filename}")

        except Exception as e:
            print(f"[{index + 1}/{len(df)}] 处理失败：{str(e)}")

    print("批量处理完成！")


if __name__ == '__main__':
    # 实际文件路径（使用原始字符串避免转义问题）
    generate_documents(
        excel_path=r'D:\测试使用\数据表.xlsx',
        template_path=r'D:\测试使用\模板.docx',
        output_dir='不动产登记文件'
    )
