import os
import copy
import re
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from docx import Document
import win32com.client as win32
from lxml import etree


def log_message(text_widget, message):
    """在日志窗口打印消息"""
    text_widget.insert(tk.END, message + "\n")
    text_widget.see(tk.END)  # 自动滚动到底部
    text_widget.update()


def get_page_count(docx_path):
    """使用 Word COM 组件统计 docx 页数"""
    word = win32.Dispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(docx_path)
    pages = doc.ComputeStatistics(2)  # 2 = wdStatisticPages
    doc.Close(False)
    word.Quit()
    return pages


def convert_doc_to_docx(doc_path):
    """将 .doc 转换为 .docx"""
    word = win32.Dispatch("Word.Application")
    word.Visible = False
    doc = word.Documents.Open(doc_path)
    new_path = doc_path + "x"  # 原名 + x 变成 .docx
    doc.SaveAs(new_path, FileFormat=16)  # 16 = wdFormatXMLDocument
    doc.Close(False)
    word.Quit()
    return new_path


def extract_area_from_filename(filename: str, custom_keywords=None) -> str:
    """提取地标名称：优先括号，否则正则匹配常见地标 + 自定义关键词"""
    # 1. 优先括号
    m = re.search(r"（(.*?)）", filename)
    if m:
        return m.group(1)

    # 2. 默认关键词
    patterns = [
        r"[\u4e00-\u9fa5]+大道\d+号?",
        r"[\u4e00-\u9fa5]+路\d+号?",
        r"[\u4e00-\u9fa5]+巷\d+号?",
        r"[\u4e00-\u9fa5]+村",
        r"[\u4e00-\u9fa5]+小区",
        r"[\u4e00-\u9fa5]+大厦",
    ]

    # 3. 动态添加自定义关键词
    if custom_keywords:
        for kw in custom_keywords:
            kw = kw.strip()
            if kw:
                patterns.append(rf"[\u4e00-\u9fa5]+{kw}\d*号?")

    # 4. 正则匹配
    for pat in patterns:
        m = re.search(pat, filename)
        if m:
            return m.group(0)

    # 5. 兜底：文件名前 15 个字
    base = os.path.splitext(filename)[0]
    return base[:15]


def replace_cover_firstpage_update(template_path, target_path, stop_page=2, custom_keywords=None):
    """替换封面+首页，并更新页数和标题"""
    template_doc = Document(template_path)
    target_doc = Document(target_path)

    # ---- 复制模板前两页 ----
    template_elems = []
    current_page = 1
    for elem in template_doc.element.body:
        template_elems.append(copy.deepcopy(elem))
        xml_str = etree.tostring(elem, encoding="unicode")
        if "w:br" in xml_str and "page" in xml_str:
            current_page += 1
            if current_page > stop_page:
                break

    # ---- 提取目标文档剩余内容（跳过前两页）----
    remaining_elems = []
    current_page = 1
    for elem in target_doc.element.body:
        xml_str = etree.tostring(elem, encoding="unicode")
        if current_page > stop_page:
            remaining_elems.append(copy.deepcopy(elem))
        if "w:br" in xml_str and "page" in xml_str:
            current_page += 1

    # ---- 重新组合文档 ----
    target_doc.element.body.clear_content()
    for elem in template_elems:
        target_doc.element.body.append(copy.deepcopy(elem))
    for elem in remaining_elems:
        target_doc.element.body.append(copy.deepcopy(elem))
    target_doc.save(target_path)

    # ---- 更新页数 ----
    total_pages = get_page_count(target_path) - 1

    # ---- 更新标题 ----
    filename = os.path.basename(target_path)
    area = extract_area_from_filename(filename, custom_keywords)
    new_title = f"{area}排水管道检测与评估技术报告"

    doc = Document(target_path)
    for para in doc.paragraphs:
        # 更新总页数
        if "报告总页数" in para.text:
            for run in para.runs:
                if re.search(r"\d+", run.text):
                    run.text = re.sub(r"\d+", str(total_pages), run.text)

        # 更新标题
        if "排水管道检测与评估技术报告" in para.text:
            for run in para.runs:
                if "排水管道检测与评估技术报告" in run.text:
                    run.text = new_title

    doc.save(target_path)
    return total_pages, new_title


# ================= GUI =================
def process_file(template_path, file_path, log_text, custom_keywords):
    try:
        # 如果是 .doc → 转换为 .docx
        if file_path.lower().endswith(".doc"):
            log_message(log_text, f"[转换] {file_path} → .docx")
            file_path = convert_doc_to_docx(file_path)

        pages, title = replace_cover_firstpage_update(template_path, file_path, custom_keywords=custom_keywords)
        log_message(log_text, f"[完成] {file_path} | 页数={pages} | 标题={title}")
    except Exception as e:
        log_message(log_text, f"[错误] {file_path} | {e}")


def process_folder(template_path, folder_path, log_text, custom_keywords):
    count = 0
    for fname in os.listdir(folder_path):
        if fname.lower().endswith((".doc", ".docx")):
            fpath = os.path.join(folder_path, fname)
            try:
                if fpath.lower().endswith(".doc"):
                    log_message(log_text, f"[转换] {fpath} → .docx")
                    fpath = convert_doc_to_docx(fpath)
                pages, title = replace_cover_firstpage_update(template_path, fpath, custom_keywords=custom_keywords)
                log_message(log_text, f"[完成] {fpath} | 页数={pages} | 标题={title}")
                count += 1
            except Exception as e:
                log_message(log_text, f"[错误] {fpath} | {e}")
    log_message(log_text, f"[总结] 共处理 {count} 个文件")


def gui_app():
    root = tk.Tk()
    root.title("批量替换报告封面 + 更新页数和标题")
    root.geometry("750x550")

    tk.Label(root, text="选择模板文件：").pack(pady=5)
    template_entry = tk.Entry(root, width=80)
    template_entry.pack(pady=5)

    def choose_template():
        path = filedialog.askopenfilename(filetypes=[("Word 文件", "*.docx")])
        if path:
            template_entry.delete(0, tk.END)
            template_entry.insert(0, path)

    tk.Button(root, text="选择模板", command=choose_template).pack(pady=5)

    # 自定义关键词输入
    tk.Label(root, text="自定义地标关键词（用逗号分隔）：").pack(pady=5)
    keyword_entry = tk.Entry(root, width=80)
    keyword_entry.pack(pady=5)
    keyword_entry.insert(0, "广场,街,城,宿舍")  # 默认值

    # 日志窗口
    log_text = scrolledtext.ScrolledText(root, width=90, height=20)
    log_text.pack(pady=10)

    def get_custom_keywords():
        return [kw.strip() for kw in keyword_entry.get().split(",") if kw.strip()]

    # 单文件处理
    def choose_file():
        template = template_entry.get()
        if not os.path.exists(template):
            messagebox.showwarning("警告", "请先选择模板文件")
            return
        fpath = filedialog.askopenfilename(filetypes=[("Word 文件", "*.doc;*.docx")])
        if fpath:
            process_file(template, fpath, log_text, get_custom_keywords())

    tk.Button(root, text="处理单个文件", command=choose_file).pack(pady=5)

    # 文件夹处理
    def choose_folder():
        template = template_entry.get()
        if not os.path.exists(template):
            messagebox.showwarning("警告", "请先选择模板文件")
            return
        folder = filedialog.askdirectory()
        if folder:
            process_folder(template, folder, log_text, get_custom_keywords())

    tk.Button(root, text="批量处理文件夹", command=choose_folder).pack(pady=5)

    root.mainloop()


if __name__ == "__main__":
    gui_app()
