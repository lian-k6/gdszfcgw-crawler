import tkinter as tk
from tkinter import ttk, filedialog, scrolledtext
import threading
import queue
import pandas as pd
from docx import Document
import os
import re
import requests
from bs4 import BeautifulSoup
import csv
import time
import random
from datetime import datetime, timedelta
import webbrowser
from docx.shared import RGBColor  # 新增导入，用于设置字体颜色
# 确保中文显示正常
import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt

plt.rcParams["font.family"] = ["SimHei", "WenQuanYi Micro Hei", "Heiti TC"]


class Application(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("多功能工具")
        self.geometry("600x400")
        self.configure(bg="#f0f0f0")

        # 创建主界面
        self.create_main_frame()

        # 队列用于线程间通信
        self.queue = queue.Queue()

        # 绑定队列处理函数
        self.process_queue()

    def create_main_frame(self):
        """创建主界面，包含两个功能按钮"""
        self.main_frame = tk.Frame(self, bg="#f0f0f0")
        self.main_frame.pack(fill=tk.BOTH, expand=True, padx=50, pady=100)

        # 标题
        title_label = tk.Label(
            self.main_frame,
            text="多功能工具",
            font=("SimHei", 24, "bold"),
            bg="#f0f0f0"
        )
        title_label.pack(pady=30)

        # 按钮框架
        button_frame = tk.Frame(self.main_frame, bg="#f0f0f0")
        button_frame.pack(fill=tk.X, pady=20)

        # 爬虫功能按钮
        spider_btn = tk.Button(
            button_frame,
            text="采购公告爬虫",
            font=("SimHei", 14),
            command=self.open_spider_frame,
            width=20,
            height=2,
            bg="#4CAF50",
            fg="white"
        )
        spider_btn.pack(side=tk.LEFT, padx=10)

        # 文档填充功能按钮
        fill_btn = tk.Button(
            button_frame,
            text="文档自动填充",
            font=("SimHei", 14),
            command=self.open_filler_frame,
            width=20,
            height=2,
            bg="#2196F3",
            fg="white"
        )
        fill_btn.pack(side=tk.RIGHT, padx=10)

    def open_spider_frame(self):
        """打开爬虫功能界面"""
        # 清空主界面
        self.clear_frame()

        # 创建爬虫界面
        self.spider_frame = tk.Frame(self, bg="#f0f0f0")
        self.spider_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        # 返回按钮
        back_btn = tk.Button(
            self.spider_frame,
            text="返回主界面",
            command=self.back_to_main,
            bg="#f44336",
            fg="white"
        )
        back_btn.pack(anchor=tk.NW, pady=(0, 10))

        # 标题
        title_label = tk.Label(
            self.spider_frame,
            text="采购公告爬虫",
            font=("SimHei", 18, "bold"),
            bg="#f0f0f0"
        )
        title_label.pack(pady=10)

        # 时间范围选择
        time_frame = tk.Frame(self.spider_frame, bg="#f0f0f0")
        time_frame.pack(fill=tk.X, pady=10)

        tk.Label(time_frame, text="爬取天数范围:", font=("SimHei", 12), bg="#f0f0f0").pack(side=tk.LEFT)
        self.day_range = tk.StringVar(value="5")
        day_entry = tk.Entry(time_frame, textvariable=self.day_range, width=5, font=("SimHei", 12))
        day_entry.pack(side=tk.LEFT, padx=10)
        tk.Label(time_frame, text="天", font=("SimHei", 12), bg="#f0f0f0").pack(side=tk.LEFT)

        # 开始按钮
        start_btn = tk.Button(
            self.spider_frame,
            text="开始爬取",
            command=self.start_spider,
            font=("SimHei", 12),
            bg="#4CAF50",
            fg="white",
            width=15,
            height=1
        )
        start_btn.pack(pady=10)

        # 进度条
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(
            self.spider_frame,
            variable=self.progress_var,
            maximum=100
        )
        self.progress_bar.pack(fill=tk.X, pady=10)

        # 状态文本框
        self.status_text = scrolledtext.ScrolledText(
            self.spider_frame,
            wrap=tk.WORD,
            font=("SimHei", 10)
        )
        self.status_text.pack(fill=tk.BOTH, expand=True, pady=10)
        self.status_text.config(state=tk.DISABLED)

    def open_filler_frame(self):
        """打开文档填充功能界面"""
        # 清空主界面
        self.clear_frame()

        # 创建文档填充界面
        self.filler_frame = tk.Frame(self, bg="#f0f0f0")
        self.filler_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        # 返回按钮
        back_btn = tk.Button(
            self.filler_frame,
            text="返回主界面",
            command=self.back_to_main,
            bg="#f44336",
            fg="white"
        )
        back_btn.pack(anchor=tk.NW, pady=(0, 10))

        # 标题
        title_label = tk.Label(
            self.filler_frame,
            text="文档自动填充",
            font=("SimHei", 18, "bold"),
            bg="#f0f0f0"
        )
        title_label.pack(pady=10)

        # Excel文件选择
        excel_frame = tk.Frame(self.filler_frame, bg="#f0f0f0")
        excel_frame.pack(fill=tk.X, pady=10)

        tk.Label(excel_frame, text="Excel文件:", font=("SimHei", 12), bg="#f0f0f0").pack(side=tk.LEFT)
        self.excel_path = tk.StringVar()
        excel_entry = tk.Entry(excel_frame, textvariable=self.excel_path, width=40, font=("SimHei", 10))
        excel_entry.pack(side=tk.LEFT, padx=10)
        excel_btn = tk.Button(excel_frame, text="浏览", command=self.select_excel)
        excel_btn.pack(side=tk.LEFT)

        # Word模板选择
        word_frame = tk.Frame(self.filler_frame, bg="#f0f0f0")
        word_frame.pack(fill=tk.X, pady=10)

        tk.Label(word_frame, text="Word模板:", font=("SimHei", 12), bg="#f0f0f0").pack(side=tk.LEFT)
        self.word_path = tk.StringVar()
        word_entry = tk.Entry(word_frame, textvariable=self.word_path, width=40, font=("SimHei", 10))
        word_entry.pack(side=tk.LEFT, padx=10)
        word_btn = tk.Button(word_frame, text="浏览", command=self.select_word)
        word_btn.pack(side=tk.LEFT)

        # 输出目录选择
        output_frame = tk.Frame(self.filler_frame, bg="#f0f0f0")
        output_frame.pack(fill=tk.X, pady=10)

        tk.Label(output_frame, text="输出目录:", font=("SimHei", 12), bg="#f0f0f0").pack(side=tk.LEFT)
        self.output_dir = tk.StringVar(value="不动产登记文件")
        output_entry = tk.Entry(output_frame, textvariable=self.output_dir, width=40, font=("SimHei", 10))
        output_entry.pack(side=tk.LEFT, padx=10)
        output_btn = tk.Button(output_frame, text="浏览", command=self.select_output_dir)
        output_btn.pack(side=tk.LEFT)

        # 记录数量选择
        count_frame = tk.Frame(self.filler_frame, bg="#f0f0f0")
        count_frame.pack(fill=tk.X, pady=10)

        tk.Label(count_frame, text="处理记录数:", font=("SimHei", 12), bg="#f0f0f0").pack(side=tk.LEFT)
        self.record_count = tk.StringVar(value="5")
        count_entry = tk.Entry(count_frame, textvariable=self.record_count, width=5, font=("SimHei", 12))
        count_entry.pack(side=tk.LEFT, padx=10)

        # 开始按钮
        start_btn = tk.Button(
            self.filler_frame,
            text="开始填充",
            command=self.start_filling,
            font=("SimHei", 12),
            bg="#2196F3",
            fg="white",
            width=15,
            height=1
        )
        start_btn.pack(pady=10)

        # 状态文本框
        self.fill_status_text = scrolledtext.ScrolledText(
            self.filler_frame,
            wrap=tk.WORD,
            font=("SimHei", 10)
        )
        self.fill_status_text.pack(fill=tk.BOTH, expand=True, pady=10)
        self.fill_status_text.config(state=tk.DISABLED)

    def back_to_main(self):
        """返回主界面"""
        self.clear_frame()
        self.create_main_frame()

    def clear_frame(self):
        """清空当前界面"""
        for widget in self.winfo_children():
            widget.destroy()

    def select_excel(self):
        """选择Excel文件"""
        file_path = filedialog.askopenfilename(
            filetypes=[("Excel files", "*.xlsx;*.xls")]
        )
        if file_path:
            self.excel_path.set(file_path)

    def select_word(self):
        """选择Word模板"""
        file_path = filedialog.askopenfilename(
            filetypes=[("Word files", "*.docx")]
        )
        if file_path:
            self.word_path.set(file_path)

    def select_output_dir(self):
        """选择输出目录"""
        dir_path = filedialog.askdirectory()
        if dir_path:
            self.output_dir.set(dir_path)

    def start_spider(self):
        """启动爬虫线程"""
        try:
            days = int(self.day_range.get())
            if days <= 0:
                self.update_spider_status("请输入有效的天数")
                return
        except ValueError:
            self.update_spider_status("请输入有效的数字")
            return

        # 清空状态
        self.status_text.config(state=tk.NORMAL)
        self.status_text.delete(1.0, tk.END)
        self.status_text.config(state=tk.DISABLED)

        # 重置进度条
        self.progress_var.set(0)

        # 在新线程中启动爬虫
        threading.Thread(
            target=self.run_spider,
            args=(days,),
            daemon=True
        ).start()

    def start_filling(self):
        """启动文档填充线程"""
        excel_path = self.excel_path.get()
        word_path = self.word_path.get()
        output_dir = self.output_dir.get()

        if not excel_path:
            self.update_fill_status("请选择Excel文件")
            return

        if not word_path:
            self.update_fill_status("请选择Word模板")
            return

        try:
            record_count = int(self.record_count.get())
            if record_count <= 0:
                self.update_fill_status("请输入有效的记录数")
                return
        except ValueError:
            self.update_fill_status("请输入有效的数字")
            return

        # 清空状态
        self.fill_status_text.config(state=tk.NORMAL)
        self.fill_status_text.delete(1.0, tk.END)
        self.fill_status_text.config(state=tk.DISABLED)

        # 在新线程中启动填充
        threading.Thread(
            target=self.run_filler,
            args=(excel_path, word_path, output_dir, record_count),
            daemon=True
        ).start()

    def update_spider_status(self, message):
        """更新爬虫状态文本框"""
        self.status_text.config(state=tk.NORMAL)
        self.status_text.insert(tk.END, message + "\n")
        self.status_text.see(tk.END)
        self.status_text.config(state=tk.DISABLED)

    def update_fill_status(self, message):
        """更新填充状态文本框"""
        self.fill_status_text.config(state=tk.NORMAL)
        self.fill_status_text.insert(tk.END, message + "\n")
        self.fill_status_text.see(tk.END)
        self.fill_status_text.config(state=tk.DISABLED)

    def process_queue(self):
        """处理队列中的消息"""
        try:
            while True:
                message = self.queue.get_nowait()
                if isinstance(message, tuple) and message[0] == "spider_progress":
                    self.progress_var.set(message[1])
                elif isinstance(message, tuple) and message[0] == "spider_status":
                    self.update_spider_status(message[1])
                elif isinstance(message, tuple) and message[0] == "fill_status":
                    self.update_fill_status(message[1])
                self.queue.task_done()
        except queue.Empty:
            pass
        self.after(100, self.process_queue)

    # 爬虫相关功能
    def run_spider(self, days):
        """运行爬虫程序"""
        # 爬虫配置
        REQUEST_DELAY_RANGE = (2, 5)
        MAX_RETRIES = 3

        USER_AGENTS = [
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.0.0 Safari/537.36",
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/115.0.0.0 Safari/537.36",
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.0.0 Safari/537.36",
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/16.5 Safari/605.1.15",
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:109.0) Gecko/20100101 Firefox/115.0",
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:102.0) Gecko/20100101 Firefox/102.0",
            "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/114.0.0.0 Safari/537.36",
            "Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:115.0) Gecko/20100101 Firefox/115.0",
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/116.0.0.0 Safari/537.36",
            "Mozilla/5.0 (iPhone; CPU iPhone OS 16_5 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/16.5 Mobile/15E148 Safari/604.1"
        ]

        def get_random_user_agent():
            return random.choice(USER_AGENTS)

        def get_random_headers():
            return {
                "User-Agent": get_random_user_agent(),
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
                "Accept-Language": "zh-CN,zh;q=0.8,en-US;q=0.5,en;q=0.3",
                "Referer": "https://ygp.gdzwfw.gov.cn/zjfwcs/gd-zjcs-pub/purchaseNotice",
                "Connection": "keep-alive",
                "Cache-Control": "max-age=0",
                "Upgrade-Insecure-Requests": "1",
            }

        def safe_request(url, method='get', data=None, headers=None, retries=0):
            try:
                if not headers:
                    headers = get_random_headers()

                delay = random.uniform(*REQUEST_DELAY_RANGE)
                time.sleep(delay)

                if method.lower() == 'post':
                    response = requests.post(url, data=data, headers=headers, timeout=10)
                else:
                    response = requests.get(url, headers=headers, timeout=10)

                response.raise_for_status()
                return response

            except requests.exceptions.RequestException as e:
                self.queue.put(("spider_status", f"请求错误: {str(e)}"))
                if retries < MAX_RETRIES:
                    self.queue.put(("spider_status", f"重试 ({retries + 1}/{MAX_RETRIES})..."))
                    new_headers = get_random_headers()
                    return safe_request(url, method, data, new_headers, retries + 1)
                self.queue.put(("spider_status", f"达到最大重试次数，请求失败"))
                return None

        def fetch_detail_page(url):
            base_url = "https://ygp.gdzwfw.gov.cn"
            full_url = base_url + url if url.startswith('/') else url

            response = safe_request(full_url)
            if not response:
                return None

            try:
                soup = BeautifulSoup(response.text, 'html.parser')
                detail = {}

                # 标题
                title_tag = soup.find('h2', class_='wrap-title-center zjcsFontFace')
                detail['标题'] = title_tag.text.strip() if title_tag else ''

                # 提取所有列表项信息
                for li in soup.find_all('li'):
                    b_tag = li.find('b')
                    if b_tag:
                        key = b_tag.text.strip().rstrip('：:')
                        value_tag = li.find('div', class_='txt') or li.find('div', class_='txt zjcsFontFace')
                        if value_tag:
                            # 清除链接和图片标签，但保留文本
                            for tag in value_tag.find_all(['a', 'img']):
                                tag.extract()
                            value = value_tag.text.strip()

                            # 映射到目标字段
                            if key == '项目业主':
                                detail['采购单位'] = value
                            elif key == '项目规模':
                                detail['预算总金额'] = value
                            elif key == '服务内容':
                                detail['项目内容'] = value
                            elif key == '采购项目名称':
                                detail['采购项目名称'] = value
                            elif key == '选取中介服务机构方式':
                                detail['选取中介服务机构方式'] = value
                            elif key == '采购项目编码':
                                detail['采购项目编码'] = value

                # 处理可能的选取方式的其他格式
                if '选取中介服务机构方式' not in detail:
                    select_mode_li = soup.find('li', string=re.compile(r'选取中介服务机构方式'))
                    if select_mode_li:
                        value_tag = select_mode_li.find('div', class_='txt') or select_mode_li.find('div',
                                                                                                    class_='txt zjcsFontFace')
                        if value_tag:
                            for img in value_tag.find_all('img'):
                                img.extract()
                            detail['选取中介服务机构方式'] = value_tag.text.strip()

                # 如果仍未找到，尝试从正文中提取
                if '选取中介服务机构方式' not in detail:
                    for p in soup.find_all('p'):
                        if '选取方式' in p.text or '直接选取' in p.text:
                            mode_text = re.search(r'(直接选取|随机抽取|竞争性谈判|询价|其他)[^。，,;；]*', p.text)
                            if mode_text:
                                detail['选取中介服务机构方式'] = mode_text.group()
                            else:
                                detail['选取中介服务机构方式'] = p.text.strip()
                            break

                # 从正文提取截止报名时间
                if '截止报名时间' not in detail:
                    content_p = soup.find('p', class_='zjcsFontFace')
                    if content_p:
                        time_match = re.search(r'\d{4}-\d{2}-\d{2} \d{2}:\d{2}', content_p.text)
                        if time_match:
                            detail['截止报名时间'] = time_match.group()

                return detail

            except Exception as e:
                self.queue.put(("spider_status", f"解析详情页 {full_url} 错误: {e}"))
                return None

        # 主爬虫逻辑
        try:
            # 计算时间范围
            end_date = datetime.now()
            start_date = end_date - timedelta(days=days)

            start_date_str = start_date.strftime('%Y-%m-%d')
            end_date_str = end_date.strftime('%Y-%m-%d')

            self.queue.put(("spider_status", f"时间范围: {start_date_str} 至 {end_date_str}"))

            # 基础请求参数
            base_params = {
                "query_params_url": "/zjfwcs/gd-zjcs-pub/purchaseNotice",
                "query_params_rest_url": "purchaseNotice/listPost",
                "reloadQueryParamsReload": "false",
                "listVo.isTrackTotalHits": "true",
                "listVo.projectName": "",
                "listVo.purOrgName": "",
                "listVo.divisionCode": "442000",
                "listVo.selectModeType": "",
                "listVo.publishDateBegin": start_date_str,
                "listVo.publishDateEnd": end_date_str,
                "listVo.projectType": "",
                "listVo.selectServiceTypes": "001",
                "pageNumber": "0"  # 初始页码
            }

            all_data = []
            seen_ids = set()
            page_number = 0
            total_processed = 0
            duplicates_removed = 0
            url = "https://ygp.gdzwfw.gov.cn/zjfwcs/gd-zjcs-pub/purchaseNotice/listPost"

            while True:
                self.queue.put(("spider_status", f"正在处理第 {page_number + 1} 页..."))

                # 更新进度条
                self.queue.put(("spider_progress", min(90, (page_number / 10) * 100)))

                # 设置当前页码
                base_params["pageNumber"] = str(page_number)

                # 发送POST请求
                response = safe_request(url, method='post', data=base_params)
                if not response:
                    self.queue.put(("spider_status", "获取列表页失败，尝试下一页..."))
                    page_number += 1
                    continue

                soup = BeautifulSoup(response.text, 'html.parser')

                # 提取表格数据
                table = soup.find('table', class_='table normal')
                if not table:
                    self.queue.put(("spider_status", "未找到表格数据，可能已无更多页面"))
                    break

                # 获取表头
                list_headers = [th.text.strip() for th in table.find('thead').find_all('th')]
                if not list_headers:
                    self.queue.put(("spider_status", "未找到表头，可能已无更多页面"))
                    break

                # 处理当前页数据
                row_count = 0
                for tr in table.find('tbody').find_all('tr'):
                    row_count += 1
                    row_data = {}

                    # 提取列表页信息
                    for i, td in enumerate(tr.find_all('td')):
                        header = list_headers[i] if i < len(list_headers) else f"字段{i}"
                        a_tag = td.find('a')
                        if a_tag:
                            row_data[header] = a_tag.text.strip()
                            row_data[f"{header}_链接"] = a_tag.get('href', '')
                            if '采购项目名称' not in row_data:
                                row_data['采购项目名称'] = a_tag.text.strip()

                    # 获取详情页数据
                    if f"{list_headers[0]}_链接" in row_data:
                        self.queue.put(("spider_status", f"  正在获取 {row_data.get(list_headers[0], '未知项目')} 的详细信息..."))
                        detail = fetch_detail_page(row_data[f"{list_headers[0]}_链接"])
                        if detail:
                            row_data.update(detail)

                            # 去重逻辑
                            unique_id = detail.get('采购项目编码')
                            if not unique_id and '标题' in detail and '采购单位' in detail:
                                unique_id = f"{detail['标题']}|{detail['采购单位']}"

                            if unique_id:
                                if unique_id not in seen_ids:
                                    seen_ids.add(unique_id)
                                    all_data.append(row_data)
                                    total_processed += 1
                                else:
                                    duplicates_removed += 1
                                    self.queue.put(("spider_status", f"  发现重复记录，已跳过 (累计移除: {duplicates_removed})"))
                            else:
                                all_data.append(row_data)
                                total_processed += 1
                                self.queue.put(("spider_status", "  无法生成唯一标识，可能存在重复记录"))

                # 如果当前页没有数据，停止分页
                if row_count == 0:
                    self.queue.put(("spider_status", "当前页无数据，停止分页"))
                    break

                # 准备下一页
                page_number += 1
                # 限制最大页数，防止无限循环
                if page_number > 20:
                    self.queue.put(("spider_status", "已达到最大页数限制，停止爬取"))
                    break

            # 定义需要保留的字段
            required_fields = [
                '标题',
                '采购项目名称',
                '预算总金额',
                '采购单位',
                '截止报名时间',
                '选取中介服务机构方式',
                '项目内容'
            ]

            # 保存为CSV文件
            filename = 'zjcs.csv'
            with open(filename, 'w', newline='', encoding='utf-8-sig') as csvfile:
                writer = csv.DictWriter(csvfile, fieldnames=required_fields)
                writer.writeheader()

                for row in all_data:
                    filtered_row = {field: row.get(field, '') for field in required_fields}
                    writer.writerow(filtered_row)

            self.queue.put(
                ("spider_status", f"处理完成，共提取 {total_processed} 条记录，移除 {duplicates_removed} 条重复记录，保存至 {filename}"))

            # 自动打开CSV文件
            try:
                if os.name == 'nt':  # Windows
                    os.startfile(filename)
                else:  # macOS/Linux
                    webbrowser.open(f'file://{os.path.abspath(filename)}')
                self.queue.put(("spider_status", f"已自动打开 {filename}"))
            except Exception as e:
                self.queue.put(("spider_status", f"自动打开文件失败: {e}，请手动打开 {filename}"))

            # 完成进度
            self.queue.put(("spider_progress", 100))

        except Exception as e:
            self.queue.put(("spider_status", f"处理数据时发生错误: {e}"))

    # 文档填充相关功能
    def run_filler(self, excel_path, template_path, output_dir, record_count):
        """运行文档填充程序"""

        def replace_text_in_doc(doc, old_text, new_text):
            """替换文档中所有出现的指定文本，保持原格式并将颜色设为黑色"""
            # 处理段落中的文本
            for para in doc.paragraphs:
                if old_text in para.text:
                    for run in para.runs:
                        if old_text in run.text:
                            # 替换文本内容
                            run.text = run.text.replace(old_text, str(new_text))
                            # 设置字体颜色为黑色
                            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色RGB值

            # 处理表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if old_text in run.text:
                                    # 替换文本内容
                                    run.text = run.text.replace(old_text, str(new_text))
                                    # 设置字体颜色为黑色
                                    run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色RGB值

        try:
            # 验证文件路径是否存在
            if not os.path.exists(excel_path):
                self.queue.put(("fill_status", f"错误：Excel文件路径不存在 - {excel_path}"))
                return
            if not os.path.exists(template_path):
                self.queue.put(("fill_status", f"错误：模板文件路径不存在 - {template_path}"))
                return

            # 创建输出目录
            os.makedirs(output_dir, exist_ok=True)
            self.queue.put(("fill_status", f"输出目录：{os.path.abspath(output_dir)}"))

            try:
                # 读取Excel数据：跳过第二行（索引1）
                df = pd.read_excel(
                    excel_path,
                    sheet_name='Sheet1',
                    header=0,  # 第一行为表头
                    skiprows=[1],  # 跳过第二行描述信息
                    nrows=record_count  # 读取指定数量的记录
                )
                self.queue.put(("fill_status", f"成功读取Excel数据，共 {len(df)} 条记录"))
                self.queue.put(("fill_status", "Excel列名：" + str(df.columns.tolist())))

            except Exception as e:
                self.queue.put(("fill_status", f"读取Excel失败：{str(e)}"))
                return

            # 字段映射（Excel列名与Word模板占位符完全一致）
            field_mapping = {
                '数据002': '数据002',
                '数据003': '数据003',
                '数据004': '数据004',
                '数据005': '数据005',
                '数据006': '数据006',
                '数据001': '数据001'  # 包含数据001用于填充和命名
            }

            # 用于命名文件的列名（数据001）
            name_column = '数据001'

            # 检查必要列是否存在
            missing_columns = [col for col in field_mapping.keys() if col not in df.columns]
            if missing_columns:
                self.queue.put(("fill_status", f"警告：Excel中缺少以下必要列，可能导致填充失败：{missing_columns}"))

            # 处理每条数据
            for index, row in df.iterrows():
                try:
                    # 打开模板
                    doc = Document(template_path)

                    # 替换所有字段
                    for excel_field, doc_placeholder in field_mapping.items():
                        if excel_field not in df.columns:
                            continue  # 跳过不存在的列
                        # 获取单元格值，空值处理为空白
                        cell_value = row[excel_field] if pd.notna(row[excel_field]) else ''
                        replace_text_in_doc(doc, doc_placeholder, cell_value)

                    # 生成文件名（使用数据001的值）
                    if name_column in df.columns and pd.notna(row[name_column]):
                        raw_name = str(row[name_column])
                        # 移除文件名中的非法字符
                        valid_name = re.sub(r'[\\/*?:"<>|]', '', raw_name)
                        filename = f"{valid_name}.docx"
                    else:
                        #  fallback：使用索引作为文件名
                        filename = f"文档_{index + 1}.docx"

                    # 保存文件
                    output_path = os.path.join(output_dir, filename)
                    doc.save(output_path)
                    self.queue.put(("fill_status", f"[{index + 1}/{len(df)}] 生成成功：{filename}"))

                except Exception as e:
                    self.queue.put(("fill_status", f"[{index + 1}/{len(df)}] 处理失败：{str(e)}"))

            self.queue.put(("fill_status", "批量处理完成！"))

        except Exception as e:
            self.queue.put(("fill_status", f"处理过程中发生错误：{str(e)}"))


if __name__ == "__main__":
    app = Application()
    app.mainloop()