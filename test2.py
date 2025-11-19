# -*- coding: utf-8 -*-
"""
optimized_yanxiantianchong_fixed.py

修正：避免在 pyqtSignal 中使用 typing 泛型导致 TypeError。
其余逻辑基于之前的优化版：解析 docx、追加写 Excel、后台线程、预览/批量处理。
依赖:
    pip install python-docx openpyxl PyQt5
"""
from __future__ import annotations

import re
import sys
import time
from typing import Dict, List, Optional, Tuple
from collections import OrderedDict
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from openpyxl.styles import Alignment

from PyQt5 import QtCore, QtWidgets

# -------------------------
# text utilities
# -------------------------
def normalize(text: Optional[str]) -> str:
    if not text:
        return ""
    return re.sub(r'\s+|[：:;；,，。.．\-–—()（）\[\]【】]', '', str(text)).lower()

def uniq_whitespace(text: Optional[str]) -> str:
    if not text:
        return ""
    return re.sub(r'\s+', ' ', text).strip()

def compress_dup_lines(text: str) -> str:
    if not text:
        return ""
    lines = [ln.rstrip() for ln in text.splitlines() if ln.rstrip() != ""]
    if not lines:
        return ""
    out = [lines[0]]
    for ln in lines[1:]:
        if ln != out[-1]:
            out.append(ln)
    return "\n".join(out)

def remove_repeated_block(text: str) -> str:
    if not text:
        return ""
    lines = [ln.rstrip() for ln in text.splitlines() if ln.strip() != ""]
    if not lines:
        return ""
    L = len(lines)
    for k in range(1, L // 2 + 1):
        if L % k != 0:
            continue
        unit = lines[:k]
        if unit * (L // k) == lines:
            return "\n".join(unit)
    return "\n".join(lines)

# -------------------------
# DocParser
# -------------------------
class DocParser:
    """Parse .docx file and extract fields. Caches by path mtime."""
    def __init__(self) -> None:
        self._cache: Dict[str, Tuple[float, OrderedDict]] = {}

    @staticmethod
    def _collect_texts_and_tables(doc: Document) -> Tuple[List[str], List[List[str]]]:
        paras = [p.text.rstrip() for p in doc.paragraphs if p.text and p.text.strip() != ""]
        table_rows: List[List[str]] = []
        for t in doc.tables:
            for r in t.rows:
                row = [c.text.strip() for c in r.cells]
                table_rows.append(row)
        return paras, table_rows

    @staticmethod
    def _parse_table_rows_to_kv(table_rows: List[List[str]]) -> OrderedDict:
        mapping: OrderedDict = OrderedDict()

        def is_label_like(token: str) -> bool:
            if not token or not token.strip():
                return False
            tok = normalize(token)
            keywords = ['建设单位','建设地点','项目名称','总面积','面积','长度','自编号','业务编号','建字','规划许可证编号','单位']
            for kw in keywords:
                if normalize(kw) in tok:
                    return True
            if token.strip().endswith('：') or token.strip().endswith(':'):
                return True
            if len(token.strip()) <= 6 and re.match(r'^[\u4e00-\u9fff]+$', token.strip()):
                return True
            return False

        for row in table_rows:
            tokens = [cell if cell is not None else "" for cell in row]
            non_empty = [t for t in tokens if t and t.strip() != ""]
            if len(non_empty) >= 2 and len(non_empty) % 2 == 0:
                ok = True
                for i in range(0, len(non_empty), 2):
                    if not is_label_like(non_empty[i]) or is_label_like(non_empty[i+1]):
                        ok = False; break
                if ok:
                    for i in range(0, len(non_empty), 2):
                        k = uniq_whitespace(non_empty[i]); v = uniq_whitespace(non_empty[i+1])
                        if k and v and k not in mapping:
                            mapping[k] = v
                    continue
            n = len(tokens); i = 0
            while i < n:
                cell = tokens[i].strip() if tokens[i] else ""
                if not cell:
                    i += 1; continue
                m = re.match(r'^(?P<k>[^：:]{1,120})[：:]\s*(?P<v>.+)$', cell)
                if m:
                    k = uniq_whitespace(m.group('k')); v = uniq_whitespace(m.group('v'))
                    if k and v and k not in mapping:
                        mapping[k] = v
                    i += 1; continue
                if is_label_like(cell):
                    j = i + 1; found = None
                    while j < n:
                        cand = tokens[j].strip()
                        if cand:
                            if not is_label_like(cand):
                                found = cand; break
                        j += 1
                    if found:
                        k = uniq_whitespace(cell); v = uniq_whitespace(found)
                        if k and v and k not in mapping:
                            mapping[k] = v
                        i = j + 1; continue
                else:
                    left_label = None
                    for l in range(i-1, -1, -1):
                        left_cell = tokens[l].strip() if tokens[l] else ""
                        if left_cell and is_label_like(left_cell):
                            left_label = left_cell; break
                    if left_label:
                        k = uniq_whitespace(left_label); v = uniq_whitespace(cell)
                        if k and v and k not in mapping:
                            mapping[k] = v
                i += 1
        cleaned = OrderedDict()
        for k, v in mapping.items():
            v2 = re.sub(r'(\b\S+\b)(\s+\1)+', r'\1', v)
            cleaned[k] = v2
        return cleaned

    @staticmethod
    def _find_inline_kv(paras: List[str], key_variants: List[str]) -> Optional[str]:
        inline_re = re.compile(r'^(?P<k>[^：:]{1,120})[：:]\s*(?P<v>.+)$')
        for p in paras:
            m = inline_re.match(p.strip())
            if m:
                k = normalize(m.group('k'))
                for kv in key_variants:
                    if kv in k:
                        return uniq_whitespace(m.group('v'))
        return None

    @staticmethod
    def _find_next_paragraph_value(paras: List[str], key_variants: List[str]) -> Optional[str]:
        for i, p in enumerate(paras):
            if any(kv in normalize(p) for kv in key_variants):
                for j in range(i+1, len(paras)):
                    if paras[j].strip():
                        return uniq_whitespace(paras[j])
        return None

    @staticmethod
    def _extract_inspection_opinion(paras: List[str], table_rows: List[List[str]]) -> str:
        for i, p in enumerate(paras):
            if '经现场验线' in p:
                block_lines = []
                for j in range(i, min(i + 12, len(paras))):
                    ln = paras[j].strip()
                    if not ln:
                        break
                    if j > i and any(k in normalize(ln) for k in ['建设单位','项目名称','建设地点','总面积','业务编号','建字','审图','核查意见']):
                        break
                    block_lines.append(ln)
                    if len(block_lines) >= 20:
                        break
                dedup = []
                for ln in block_lines:
                    if not dedup or ln != dedup[-1]:
                        dedup.append(ln)
                return "\n".join(dedup).strip()
        norm_paras = [re.sub(r'\s+', '', p) for p in paras]
        label = '核查意见'
        for i in range(len(norm_paras)):
            accum = ''
            for k in range(0, 4):
                if i + k < len(norm_paras):
                    accum += norm_paras[i + k]
                    if label in accum:
                        start = i + k + 1
                        block_lines = []
                        for j in range(start, min(start + 12, len(paras))):
                            line = paras[j].strip()
                            if not line:
                                break
                            if any(h in normalize(line) for h in ['建设单位','项目名称','建设地点','总面积','业务编号','建字','审图']):
                                break
                            block_lines.append(line)
                            if len(block_lines) >= 20:
                                break
                        dedup = []
                        for ln in block_lines:
                            if not dedup or ln != dedup[-1]:
                                dedup.append(ln)
                        return "\n".join(dedup).strip()
        for row in table_rows:
            for idx, c in enumerate(row):
                if c and '经现场验线' in c:
                    parts = [c.strip()]
                    if idx + 1 < len(row) and row[idx+1].strip() and not any(h in normalize(row[idx+1]) for h in ['建设单位','项目名称','建设地点','总面积','业务编号','建字']):
                        parts.append(row[idx+1].strip())
                    dedup = []
                    for ln in "\n".join(parts).splitlines():
                        if not dedup or ln.strip() != dedup[-1].strip():
                            dedup.append(ln.rstrip())
                    return "\n".join(dedup).strip()
        return ""

    @staticmethod
    def _extract_build_number(paras: List[str], table_rows: List[List[str]]) -> str:
        joined = "\n".join(paras) + "\n" + "\n".join([" ".join([c for c in row if c and c.strip()]) for row in table_rows])
        m = re.search(r'建字第[\s\u3000]*([^号\n]{1,40}?)\s*号', joined)
        if m:
            token = re.sub(r'[^\dA-Za-z\u4e00-\u9fff\-\_]+', '', m.group(1).strip())
            if token:
                return f"建字第{token}号"
        for p in paras:
            m2 = re.search(r'建字第[\s\u3000]*([^号\n]{1,40}?)\s*号', p)
            if m2:
                token = re.sub(r'[^\dA-Za-z\u4e00-\u9fff\-\_]+', '', m2.group(1).strip())
                return f"建字第{token}号"
        for row in table_rows:
            for c in row:
                if not c or c.strip() == "":
                    continue
                m3 = re.search(r'建字第[\s\u3000]*([^号\n]{1,40}?)\s*号', c)
                if m3:
                    token = re.sub(r'[^\dA-Za-z\u4e00-\u9fff\-\_]+', '', m3.group(1).strip())
                    return f"建字第{token}号"
        return ""

    def parse(self, path: Path) -> OrderedDict:
        path = path.resolve()
        if not path.exists():
            raise FileNotFoundError(str(path))
        mtime = path.stat().st_mtime
        key = str(path)
        cached = self._cache.get(key)
        if cached and cached[0] == mtime:
            return cached[1].copy()
        doc = Document(str(path))
        paras, table_rows = self._collect_texts_and_tables(doc)
        data: OrderedDict = OrderedDict()

        joined = "\n".join([p for p in paras if p and p.strip()]) + "\n" + "\n".join([" ".join([c for c in row if c and c.strip()]) for row in table_rows])

        m = re.search(r'自编号[：:\s]*([A-Za-z0-9\-]+)', joined)
        data['自编号'] = m.group(1).strip() if m else (self._find_inline_kv(paras, ['自编号']) or "")

        bizs = re.findall(r'业务编号[：:\s]*([0-9A-Za-z\-]+)', joined)
        first_biz_line: Optional[str] = None
        for p in paras:
            if '业务编号' in p:
                first_biz_line = p.strip(); break
        if not first_biz_line:
            for row in table_rows:
                for c in row:
                    if '业务编号' in c:
                        first_biz_line = c.strip(); break
                if first_biz_line: break

        if bizs:
            data['规划许可证附件的项目编号+告知书日期（复制报告业务编号）'] = bizs[0]
        else:
            if first_biz_line:
                mnum = re.search(r'([0-9A-Za-z\-]{3,})', first_biz_line)
                data['规划许可证附件的项目编号+告知书日期（复制报告业务编号）'] = mnum.group(1) if mnum else first_biz_line
            else:
                data['规划许可证附件的项目编号+告知书日期（复制报告业务编号）'] = ""

        data['规划许可证业务编号（复制报告，对照公告牌'] = bizs[1] if len(bizs) >= 2 else ""
        data['业务编号'] = bizs[0] if bizs else ""
        data['建字（复制报告，看许可证-对照文件名'] = self._extract_build_number(paras, table_rows)

        table_map = self._parse_table_rows_to_kv(table_rows)

        unit = None
        for k, v in table_map.items():
            if '建设单位' in normalize(k) or '单位（盖章）' in normalize(k) or normalize(k) == '单位':
                if v and v.strip():
                    unit = v; break
        if not unit:
            unit = self._find_inline_kv(paras, ['建设单位','单位（盖章）','单位'])
        if not unit:
            unit = self._find_next_paragraph_value(paras, ['建设单位','单位'])
        if not unit:
            munit = re.search(r'建设单位[：:\s]*([^\n,，。]+)', joined)
            if munit:
                unit = munit.group(1).strip()
        data['建设单位'] = uniq_whitespace(unit) if unit else ""

        loc = None
        for k, v in table_map.items():
            if '建设地点' in normalize(k) or normalize(k) == '地点':
                if v and v.strip():
                    loc = v; break
        if not loc:
            loc = self._find_inline_kv(paras, ['建设地点','地点'])
        if not loc:
            loc = self._find_next_paragraph_value(paras, ['建设地点','地点'])
        if not loc:
            mloc = re.search(r'建设地点[：:\s]*([^\n,，。]+)', joined)
            if mloc:
                loc = mloc.group(1).strip()
        data['建设地点'] = uniq_whitespace(loc) if loc else ""

        proj = None
        for k, v in table_map.items():
            if '项目名称' in normalize(k) or ('项目' in normalize(k) and '名称' in normalize(k)):
                if v and v.strip():
                    proj = v; break
        if not proj:
            proj = self._find_inline_kv(paras, ['项目名称','项目'])
        if not proj:
            proj = self._find_next_paragraph_value(paras, ['项目名称','项目'])
        if not proj:
            mproj = re.search(r'项目名称[：:\s]*([^\n,，。]+)', joined)
            if mproj:
                proj = mproj.group(1).strip()
        data['项目名称'] = uniq_whitespace(proj) if proj else ""

        area = None
        for k, v in table_map.items():
            if '总面积' in normalize(k) or '面积' in normalize(k) or '长度' in normalize(k):
                if v and v.strip():
                    area = v; break
        if not area:
            area = self._find_inline_kv(paras, ['总面积','面积','长度'])
        if not area:
            area = self._find_next_paragraph_value(paras, ['总面积','面积','长度'])
        if area:
            mnum = re.search(r'([0-9.]+)', area)
            area = mnum.group(1) if mnum else area
        data['总面积/长度'] = area or ""

        opinion = self._extract_inspection_opinion(paras, table_rows)
        data['报告里文字说明'] = opinion or ""

        self._cache[key] = (mtime, data.copy())
        return data

# -------------------------
# ExcelWriter
# -------------------------
class ExcelWriter:
    def __init__(self, excel_path: Path, sheet_name: Optional[str] = None) -> None:
        self.excel_path = excel_path.resolve()
        if not self.excel_path.exists():
            raise FileNotFoundError(f"Excel file not found: {self.excel_path}")
        self.wb = load_workbook(str(self.excel_path))
        self.ws = self.wb[sheet_name] if sheet_name and sheet_name in self.wb.sheetnames else self.wb.active
        self.header_cells = [c.value if c.value is not None else "" for c in self.ws[1]]
        self.header_map: Dict[str, int] = {str(h).strip(): idx + 1 for idx, h in enumerate(self.header_cells) if h and str(h).strip()}

    def add_row(self, data: OrderedDict, source_filename: Optional[str] = None) -> Tuple[int, List[str]]:
        target_row = self.ws.max_row + 1
        written = 0; skipped: List[str] = []
        for k, v in data.items():
            if k in self.header_map:
                cell = self.ws.cell(row=target_row, column=self.header_map[k])
                cell.value = v
                if isinstance(v, str) and '\n' in v:
                    cell.alignment = Alignment(wrapText=True)
                written += 1
            else:
                skipped.append(k)
        for candidate in ("SourceFile", "文件名", "来源文件", "来源"):
            if candidate in self.header_map:
                self.ws.cell(row=target_row, column=self.header_map[candidate]).value = Path(source_filename).name if source_filename else ""
                written += 1
                break
        return written, skipped

    def save(self) -> None:
        self.wb.save(str(self.excel_path))
        time.sleep(0.1)

# -------------------------
# Workers (signals use runtime types)
# -------------------------
class BatchWorker(QtCore.QThread):
    progress = QtCore.pyqtSignal(int, int)
    log = QtCore.pyqtSignal(str)
    finished = QtCore.pyqtSignal(int, int)

    def __init__(self, files: List[Path], excel_path: Path, sheet_name: Optional[str] = None) -> None:
        super().__init__()
        self.files = [p.resolve() for p in files]
        self.excel_path = excel_path.resolve()
        self.sheet_name = sheet_name
        self._running = True
        self._parser = DocParser()

    def run(self) -> None:
        total = len(self.files)
        processed = 0
        try:
            writer = ExcelWriter(self.excel_path, self.sheet_name)
        except Exception as e:
            self.log.emit(f"[错误] 打开 Excel 失败: {e}")
            self.finished.emit(processed, total)
            return
        try:
            for p in self.files:
                if not self._running:
                    self.log.emit("[停止] 收到停止请求")
                    break
                if not p.exists():
                    self.log.emit(f"[跳过] 文件不存在: {p}")
                    processed += 1
                    self.progress.emit(processed, total)
                    continue
                try:
                    self.log.emit(f"[解析] {p.name}")
                    data = self._parser.parse(p)
                    written, skipped = writer.add_row(data, source_filename=str(p))
                    self.log.emit(f"[{p.name}] 写入 {written} 列，跳过 {len(skipped)} 列。")
                    if skipped:
                        self.log.emit("  跳过列: " + ", ".join(skipped))
                except Exception as e:
                    self.log.emit(f"[错误] 处理 {p.name} 失败: {e}")
                processed += 1
                self.progress.emit(processed, total)
            try:
                writer.save()
                self.log.emit("[完成] Excel 保存成功")
            except Exception as e:
                self.log.emit(f"[错误] Excel 保存失败: {e}")
        finally:
            self.finished.emit(processed, total)

    def stop(self) -> None:
        self._running = False

class PreviewWorker(QtCore.QThread):
    # Use object to avoid typing generics in signal declaration
    done = QtCore.pyqtSignal(object)
    error = QtCore.pyqtSignal(str)
    log = QtCore.pyqtSignal(str)

    def __init__(self, path: Path) -> None:
        super().__init__()
        self.path = path.resolve()
        self._parser = DocParser()

    def run(self) -> None:
        if not self.path.exists():
            self.error.emit("文件不存在: " + str(self.path)); return
        try:
            self.log.emit(f"[预览] 解析 {self.path.name} ...")
            data = self._parser.parse(self.path)
            self.done.emit(data)  # emit OrderedDict (object)
        except Exception as e:
            self.error.emit(str(e))

# -------------------------
# GUI
# -------------------------
class MainWindow(QtWidgets.QWidget):
    def __init__(self) -> None:
        super().__init__()
        self.setWindowTitle("Word -> Excel 提取（修正版）")
        self.resize(980, 640)
        self._worker: Optional[BatchWorker] = None
        self._preview: Optional[PreviewWorker] = None
        self._build_ui()

    def _build_ui(self) -> None:
        layout = QtWidgets.QVBoxLayout(self)
        title = QtWidgets.QLabel("Word → Excel 批量提取（修正版）")
        title.setStyleSheet("font-size:16px; font-weight:bold;"); title.setAlignment(QtCore.Qt.AlignCenter)
        layout.addWidget(title)

        form = QtWidgets.QGridLayout()
        self.word_edit = QtWidgets.QLineEdit(); btn_words = QtWidgets.QPushButton("选择 Word(.docx) 多个")
        btn_words.clicked.connect(self.choose_words)
        self.excel_edit = QtWidgets.QLineEdit(); btn_excel = QtWidgets.QPushButton("选择 Excel(.xlsx)")
        btn_excel.clicked.connect(self.choose_excel)
        self.sheet_edit = QtWidgets.QLineEdit(); self.sheet_edit.setPlaceholderText("Sheet 名 (留空=active)")
        form.addWidget(QtWidgets.QLabel("Word 文件:"), 0, 0)
        form.addWidget(self.word_edit, 0, 1); form.addWidget(btn_words, 0, 2)
        form.addWidget(QtWidgets.QLabel("Excel 文件:"), 1, 0)
        form.addWidget(self.excel_edit, 1, 1); form.addWidget(btn_excel, 1, 2)
        form.addWidget(QtWidgets.QLabel("Sheet 名:"), 2, 0); form.addWidget(self.sheet_edit, 2, 1)
        layout.addLayout(form)

        h = QtWidgets.QHBoxLayout()
        self.preview_btn = QtWidgets.QPushButton("预览（第一个）"); self.preview_btn.clicked.connect(self.preview_first)
        self.start_btn = QtWidgets.QPushButton("开始写入"); self.start_btn.clicked.connect(self.start_batch)
        self.stop_btn = QtWidgets.QPushButton("停止"); self.stop_btn.clicked.connect(self.stop_batch); self.stop_btn.setEnabled(False)
        h.addWidget(self.preview_btn); h.addWidget(self.start_btn); h.addWidget(self.stop_btn); h.addStretch()
        layout.addLayout(h)

        self.progress = QtWidgets.QProgressBar(); layout.addWidget(self.progress)
        self.log_area = QtWidgets.QPlainTextEdit(); self.log_area.setReadOnly(True); layout.addWidget(self.log_area, 1)

    def log(self, msg: str) -> None:
        timestamp = time.strftime("%H:%M:%S")
        self.log_area.appendPlainText(f"[{timestamp}] {msg}")

    def choose_words(self) -> None:
        files, _ = QtWidgets.QFileDialog.getOpenFileNames(self, "选择 Word 文档", "", "Word 文档 (*.docx)")
        if files:
            self.word_edit.setText(";".join(files))

    def choose_excel(self) -> None:
        f, _ = QtWidgets.QFileDialog.getOpenFileName(self, "选择 Excel 文件", "", "Excel 文件 (*.xlsx)")
        if f:
            self.excel_edit.setText(f)

    def preview_first(self) -> None:
        files = [p for p in self.word_edit.text().split(";") if p.strip()]
        if not files:
            self.log("请先选择 Word 文件用于预览。"); return
        first = Path(files[0]).resolve()
        if not first.exists():
            self.log("第一个文件不存在: " + str(first)); return
        self.preview_btn.setEnabled(False)
        self._preview = PreviewWorker(first)
        self._preview.log.connect(self.log)
        self._preview.done.connect(self.on_preview_done)
        self._preview.error.connect(self.on_preview_error)
        self._preview.start()
        self.log(f"[预览] 开始解析 {first.name}")

    def on_preview_done(self, data_obj: object) -> None:
        self.preview_btn.setEnabled(True)
        self.log("[预览] 解析完成，结果如下：")
        if isinstance(data_obj, OrderedDict):
            items = list(data_obj.items())
        elif isinstance(data_obj, dict):
            items = list(data_obj.items())
        else:
            # best-effort print
            self.log(f"返回类型: {type(data_obj)}, 内容: {data_obj}")
            return
        for k, v in items:
            self.log(f"  {k} -> {v}")

    def on_preview_error(self, msg: str) -> None:
        self.preview_btn.setEnabled(True)
        self.log("[预览错误] " + msg)

    def start_batch(self) -> None:
        files = [Path(p).resolve() for p in self.word_edit.text().split(";") if p.strip()]
        excel_file = Path(self.excel_edit.text().strip()) if self.excel_edit.text().strip() else None
        sheet = self.sheet_edit.text().strip() or None
        if not files:
            self.log("请选择 Word 文件"); return
        if not excel_file or not excel_file.exists():
            self.log("请选择存在的 Excel 文件"); return
        self.start_btn.setEnabled(False); self.preview_btn.setEnabled(False); self.stop_btn.setEnabled(True); self.progress.setValue(0)
        self._worker = BatchWorker(files, excel_file, sheet)
        self._worker.log.connect(self.log)
        self._worker.progress.connect(lambda a,b: self.progress.setValue(int(a*100/b) if b else 0))
        self._worker.finished.connect(self.on_finished)
        self._worker.start()
        self.log("[批量] 开始处理")

    def stop_batch(self) -> None:
        if self._worker:
            self._worker.stop(); self.log("已请求停止（当前文件处理完后停止）")

    def on_finished(self, processed: int, total: int) -> None:
        self.log(f"[批量] 完成: {processed}/{total}")
        self.start_btn.setEnabled(True); self.preview_btn.setEnabled(True); self.stop_btn.setEnabled(False)
        if total:
            self.progress.setValue(100 if processed == total else int(processed*100/total))

# -------------------------
# Entrypoint
# -------------------------
def main() -> None:
    app = QtWidgets.QApplication(sys.argv)
    win = MainWindow(); win.show()
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()
