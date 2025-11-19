# -*- coding: utf-8 -*-
"""
yanxiantianchong.py

在原始基础上做两项修复：
1) 规划许可证附件的项目编号+告知书日期 -- 只保留业务编号（优先 bizs[0]）
2) 报告里文字说明（核查意见） -- 增强去重，删除整段重复和相邻重复行

其余逻辑（表格解析、写 Excel、GUI）保持不改。
依赖: python-docx, openpyxl, PyQt5
"""

import sys
import os
import re
from collections import OrderedDict
from docx import Document
from openpyxl import load_workbook
from openpyxl.styles import Alignment
from PyQt5 import QtWidgets, QtCore

# -------------------------
# helpers
# -------------------------
def normalize(s):
    if not s:
        return ""
    return re.sub(r'\s+|[：:;；,，。.．\-–—()（）\[\]【】]', '', str(s)).lower()

def uniq_whitespace(s):
    return re.sub(r'\s+', ' ', s.strip()) if isinstance(s, str) else s

def compress_dup_lines(text):
    """去除相邻重复行并返回合并字符串"""
    if not text:
        return text
    lines = [ln.rstrip() for ln in text.splitlines() if ln.rstrip() != ""]
    if not lines:
        return ""
    out = [lines[0]]
    for ln in lines[1:]:
        if ln != out[-1]:
            out.append(ln)
    return "\n".join(out)

def remove_repeated_block(text):
    """
    如果 text 由若干行组成，且这些行是某个最小块重复多次（例如 [A,B,A,B]），
    则返回该最小块（'A\nB'），否则返回原始去除空行的文本。
    这样可以去掉整段被重复粘贴两次/多次的情况。
    """
    if not text:
        return text
    # 保持原行顺序但去掉全空行
    lines = [ln.rstrip() for ln in text.splitlines() if ln.strip() != ""]
    if not lines:
        return ""
    L = len(lines)
    # 尝试找到最小重复单元
    for k in range(1, L // 2 + 1):
        if L % k != 0:
            continue
        unit = lines[:k]
        if unit * (L // k) == lines:
            return "\n".join(unit)
    # 额外处理：如果整段文本重复自身两次（可能因为不同的空白或轻微差异），
    # 也尝试用字符串方式检测完全重复的双份（宽松匹配空格）
    joined = "\n".join(lines)
    # escape for regex, allow arbitrary whitespace between repeats
    esc = re.escape(joined)
    pattern = r'^(?:\s*' + esc + r'\s*)+$'
    if re.match(pattern, text.strip(), flags=re.DOTALL):
        return joined
    return joined

# -------------------------
# collect doc texts
# -------------------------
def collect_doc_texts(doc):
    paras = [p.text.rstrip() for p in doc.paragraphs if p.text and p.text.strip() != ""]
    table_rows = []
    for t in doc.tables:
        for r in t.rows:
            row = [c.text.strip() for c in r.cells]
            table_rows.append(row)
    return paras, table_rows

# -------------------------
# robust建字提取（保持原样）
# -------------------------
def extract_build_number(paras, table_rows):
    joined = "\n".join(paras) + "\n" + "\n".join([" ".join([c for c in row if c and c.strip()]) for row in table_rows])
    m = re.search(r'建字第[\s\u3000]*([^号\n]{1,40}?)\s*号', joined)
    if m:
        token = m.group(1).strip()
        token_clean = re.sub(r'[^\dA-Za-z\u4e00-\u9fff\-\_]+', '', token)
        if token_clean:
            return f"建字第{token_clean}号"
    for p in paras:
        m2 = re.search(r'建字第[\s\u3000]*([^号\n]{1,40}?)\s*号', p)
        if m2:
            token = m2.group(1).strip()
            token = re.sub(r'[^\dA-Za-z\u4e00-\u9fff\-\_]+','', token)
            return f"建字第{token}号"
    for row in table_rows:
        for c in row:
            if not c or c.strip() == "":
                continue
            m3 = re.search(r'建字第[\s\u3000]*([^号\n]{1,40}?)\s*号', c)
            if m3:
                token = m3.group(1).strip()
                token = re.sub(r'[^\dA-Za-z\u4e00-\u9fff\-\_]+','', token)
                return f"建字第{token}号"
    return ""

# -------------------------
# 改进后的核查意见提取（使用 remove_repeated_block + compress_dup_lines）
# -------------------------
def extract_inspection_opinion(paras, table_rows):
    """
    优先:
      1) 在段落中查找 '经现场验线' 并收集该段 + 后续若干相关段（直到下一个字段或空行）
      2) 若没有, 查找 '核查意见' 标签（支持拆分）
      3) 若仍没有, 在表格中查找 '经现场验线' 单元格作为备用
    对找到的内容:
      - 先用 remove_repeated_block() 去掉整段重复（如整段重复两次）
      - 再用 compress_dup_lines() 去掉相邻重复行
      - 最终返回一块唯一文本
    """
    # 1) 段落中找 '经现场验线'
    for idx, p in enumerate(paras):
        if '经现场验线' in p:
            collected = []
            for j in range(idx, min(idx + 12, len(paras))):
                txt = paras[j].strip()
                if not txt:
                    break
                # 如果遇到明显的新字段/标题则停止
                if j > idx and any(k in normalize(txt) for k in ['建设单位','项目名称','建设地点','总面积','业务编号','建字','审图','核查意见']):
                    break
                collected.append(txt)
            if collected:
                block = "\n".join(collected).strip()
                # remove repeated block sequences (e.g., whole block pasted twice)
                block = remove_repeated_block(block)
                # remove adjacent duplicate lines inside block
                block = compress_dup_lines(block)
                return block

    # 2) 段落中找 '核查意见'（允许拆分）
    norm_paras = [re.sub(r'\s+', '', p) for p in paras]
    label = '核查意见'
    for i in range(len(norm_paras)):
        accum = ''
        for k in range(4):
            if i + k < len(norm_paras):
                accum += norm_paras[i + k]
                if label in accum:
                    start_idx = i + k + 1
                    collected = []
                    for j in range(start_idx, min(start_idx + 12, len(paras))):
                        txt = paras[j].strip()
                        if not txt:
                            break
                        if any(h in normalize(txt) for h in ['建设单位','项目名称','建设地点','总面积','业务编号','建字','审图']):
                            break
                        collected.append(txt)
                    if collected:
                        block = "\n".join(collected).strip()
                        block = remove_repeated_block(block)
                        block = compress_dup_lines(block)
                        return block
                    else:
                        return ""

    # 3) 表格中找 '经现场验线'（仅当段落中未找到）
    for row in table_rows:
        for idx, c in enumerate(row):
            if c and '经现场验线' in c:
                parts = [c.strip()]
                if idx + 1 < len(row) and row[idx+1].strip() and not any(h in normalize(row[idx+1]) for h in ['建设单位','项目名称','建设地点','总面积','业务编号','建字']):
                    parts.append(row[idx+1].strip())
                block = "\n".join(parts).strip()
                block = remove_repeated_block(block)
                block = compress_dup_lines(block)
                return block
    return ""

# -------------------------
# other extraction helpers (保持原样)
# -------------------------
def find_inline_kv(paras, keys):
    inline_re = re.compile(r'^(?P<k>[^：:]{1,120})[：:]\s*(?P<v>.+)$')
    for p in paras:
        m = inline_re.match(p.strip())
        if m:
            k = normalize(m.group('k'))
            for kw in keys:
                if normalize(kw) in k:
                    return uniq_whitespace(m.group('v'))
    return None

def find_next_paragraph_value(paras, keys):
    for i, p in enumerate(paras):
        if any(normalize(k) in normalize(p) for k in keys):
            for j in range(i+1, len(paras)):
                if paras[j].strip():
                    return uniq_whitespace(paras[j])
    return None

def extract_field_from_table_by_keywords(table_rows, keywords):
    header_keywords = ['建设单位','建设地点','项目名称','总面积','面积','长度','业务编号','建字','审图','自编号']
    for row in table_rows:
        n = len(row)
        for i, cell in enumerate(row):
            if not cell or not cell.strip():
                continue
            if any(normalize(k) in normalize(cell) for k in keywords):
                j = i + 1
                while j < n:
                    cand = row[j].strip()
                    if cand:
                        if any(normalize(hk) in normalize(cand) for hk in header_keywords):
                            j += 1
                            continue
                        return uniq_whitespace(cand)
                    j += 1
                m = re.search(r'[：:]\s*(.+)$', cell)
                if m:
                    return uniq_whitespace(m.group(1))
    return None

# -------------------------
# main extraction combining above components (仅此处修改两个逻辑点)
# -------------------------
def extract_fields_from_docx(path):
    doc = Document(path)
    paras, table_rows = collect_doc_texts(doc)
    joined = "\n".join([p for p in paras if p and p.strip()]) + "\n" + "\n".join([" ".join([c for c in row if c and c.strip()]) for row in table_rows])

    data = OrderedDict()

    # 自编号
    m = re.search(r'自编号[：:\s]*([A-Za-z0-9\-]+)', joined)
    data['自编号'] = m.group(1).strip() if m else (find_inline_kv(paras, ['自编号']) or "")

    # 业务编号(s)
    bizs = re.findall(r'业务编号[：:\s]*([0-9A-Za-z\-]+)', joined)
    # first_biz_line (整行) - 仍保留用于备选，但主字段改为只写数字
    first_biz_line = None
    for p in paras:
        if '业务编号' in p:
            first_biz_line = p.strip()
            break
    if not first_biz_line:
        for row in table_rows:
            for c in row:
                if '业务编号' in c:
                    first_biz_line = c.strip(); break
            if first_biz_line:
                break

    # ---------- 修改点1: 规划许可证附件的项目编号+告知书日期 只保留业务编号 ----------
    if bizs:
        data['规划许可证附件的项目编号+告知书日期（复制报告业务编号）'] = bizs[0]
    else:
        # 从整行提取数字/字母序列作为备选
        if first_biz_line:
            mnum = re.search(r'([0-9A-Za-z\-]{3,})', first_biz_line)
            data['规划许可证附件的项目编号+告知书日期（复制报告业务编号）'] = mnum.group(1) if mnum else first_biz_line
        else:
            data['规划许可证附件的项目编号+告知书日期（复制报告业务编号）'] = ""

    data['规划许可证业务编号（复制报告，对照公告牌'] = bizs[1] if len(bizs) >= 2 else ""
    data['业务编号'] = bizs[0] if bizs else ""

    # 建字
    data['建字（复制报告，看许可证-对照文件名'] = extract_build_number(paras, table_rows)

    # 建设单位
    unit = extract_field_from_table_by_keywords(table_rows, ['建设单位', '单位（盖章）', '单位'])
    if not unit:
        unit = find_inline_kv(paras, ['建设单位', '单位（盖章）', '单位'])
    if not unit:
        unit = find_next_paragraph_value(paras, ['建设单位','单位'])
    if not unit:
        munit = re.search(r'建设单位[：:\s]*([^\n,，。]+)', joined)
        if munit:
            unit = munit.group(1).strip()
    data['建设单位'] = uniq_whitespace(unit) if unit else ""

    # 建设地点
    loc = extract_field_from_table_by_keywords(table_rows, ['建设地点', '地点'])
    if not loc:
        loc = find_inline_kv(paras, ['建设地点', '地点'])
    if not loc:
        loc = find_next_paragraph_value(paras, ['建设地点','地点'])
    if not loc:
        mloc = re.search(r'建设地点[：:\s]*([^\n,，。]+)', joined)
        if mloc:
            loc = mloc.group(1).strip()
    data['建设地点'] = uniq_whitespace(loc) if loc else ""

    # 项目名称
    proj = extract_field_from_table_by_keywords(table_rows, ['项目名称', '项目', '名称'])
    if not proj:
        proj = find_inline_kv(paras, ['项目名称', '项目'])
    if not proj:
        proj = find_next_paragraph_value(paras, ['项目名称', '项目'])
    if not proj:
        mproj = re.search(r'项目名称[：:\s]*([^\n,，。]+)', joined)
        if mproj:
            proj = mproj.group(1).strip()
    data['项目名称'] = uniq_whitespace(proj) if proj else ""

    # 总面积/长度
    area = extract_field_from_table_by_keywords(table_rows, ['总面积', '面积', '长度'])
    if not area:
        area = find_inline_kv(paras, ['总面积','面积','长度'])
    if not area:
        area = find_next_paragraph_value(paras, ['总面积','面积','长度'])
    if area:
        m2 = re.search(r'([0-9.]+)', area)
        area = m2.group(1) if m2 else area
    data['总面积/长度'] = area or ""

    # 报告里文字说明 -> 使用改进后的函数（先 remove_repeated_block 再 compress_dup_lines）
    opinion = extract_inspection_opinion(paras, table_rows)
    if opinion:
        # 先整体去重（处理整段被重复粘贴），再去相邻重复行
        opinion = remove_repeated_block(opinion)
        opinion = compress_dup_lines(opinion)
    data['报告里文字说明'] = opinion or ""

    return data

# -------------------------
# Excel 写入（不新增列）
# -------------------------
def write_to_excel_no_add(excel_path, sheet_name, data_row, source_filename=None):
    wb = load_workbook(excel_path)
    ws = wb[sheet_name] if sheet_name and sheet_name in wb.sheetnames else wb.active
    header_cells = [c.value if c.value is not None else "" for c in ws[1]]
    header_map = {str(h).strip(): idx+1 for idx, h in enumerate(header_cells) if h and str(h).strip()}
    target_row = ws.max_row + 1
    written = 0
    skipped = []
    for k, v in data_row.items():
        if k in header_map:
            ws.cell(row=target_row, column=header_map[k]).value = v
            if isinstance(v, str) and '\n' in v:
                ws.cell(row=target_row, column=header_map[k]).alignment = Alignment(wrapText=True)
            written += 1
        else:
            skipped.append(k)
    for candidate in ("SourceFile", "文件名", "来源文件", "来源"):
        if candidate in header_map:
            ws.cell(row=target_row, column=header_map[candidate]).value = os.path.basename(source_filename) if source_filename else ""
            written += 1
            break
    wb.save(excel_path)
    return written, skipped

# -------------------------
# Worker & GUI (保持原样)
# -------------------------
class Worker(QtCore.QThread):
    progress = QtCore.pyqtSignal(int,int)
    log = QtCore.pyqtSignal(str)
    finished = QtCore.pyqtSignal(int,int)
    def __init__(self, files, excel, sheet):
        super().__init__()
        self.files = files; self.excel = excel; self.sheet = sheet; self._running = True
    def run(self):
        total = len(self.files); processed = 0
        for f in self.files:
            if not self._running:
                break
            if not os.path.exists(f):
                self.log.emit(f"[跳过] 文件不存在: {f}"); self.progress.emit(processed, total); continue
            try:
                data = extract_fields_from_docx(f)
                written, skipped = write_to_excel_no_add(self.excel, self.sheet, data, source_filename=f)
                processed += 1
                self.log.emit(f"[{os.path.basename(f)}] 写入 {written} 列，跳过 {len(skipped)} 列。")
                if skipped:
                    self.log.emit("  跳过列: " + ", ".join(skipped))
            except Exception as e:
                self.log.emit(f"[错误] 处理 {os.path.basename(f)} 失败: {e}")
            self.progress.emit(processed, total)
        self.finished.emit(processed, total)
    def stop(self): self._running = False

class MainWindow(QtWidgets.QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Word->Excel 提取（建字与核查意见修正版）")
        self.resize(980, 640)
        self.worker = None
        self._build_ui()
    def _build_ui(self):
        L = QtWidgets.QVBoxLayout(self)
        grid = QtWidgets.QGridLayout()
        self.word_edit = QtWidgets.QLineEdit(); btn_w = QtWidgets.QPushButton("选择 Word 多个"); btn_w.clicked.connect(self.select_words)
        self.excel_edit = QtWidgets.QLineEdit(); btn_x = QtWidgets.QPushButton("选择 Excel"); btn_x.clicked.connect(self.select_excel)
        self.sheet_edit = QtWidgets.QLineEdit()
        grid.addWidget(QtWidgets.QLabel("Word 文件:"),0,0); grid.addWidget(self.word_edit,0,1); grid.addWidget(btn_w,0,2)
        grid.addWidget(QtWidgets.QLabel("Excel 文件:"),1,0); grid.addWidget(self.excel_edit,1,1); grid.addWidget(btn_x,1,2)
        grid.addWidget(QtWidgets.QLabel("Sheet 名称(留空=active):"),2,0); grid.addWidget(self.sheet_edit,2,1)
        L.addLayout(grid)
        h = QtWidgets.QHBoxLayout()
        self.preview_btn = QtWidgets.QPushButton("预览(第一个)"); self.preview_btn.clicked.connect(self.preview_first)
        self.start_btn = QtWidgets.QPushButton("开始写入"); self.start_btn.clicked.connect(self.start_processing)
        self.stop_btn = QtWidgets.QPushButton("停止"); self.stop_btn.clicked.connect(self.stop_processing); self.stop_btn.setEnabled(False)
        h.addWidget(self.preview_btn); h.addWidget(self.start_btn); h.addWidget(self.stop_btn); h.addStretch()
        L.addLayout(h)
        self.progress = QtWidgets.QProgressBar(); L.addWidget(self.progress)
        self.log = QtWidgets.QPlainTextEdit(); self.log.setReadOnly(True); L.addWidget(self.log,1)
    def select_words(self):
        files, _ = QtWidgets.QFileDialog.getOpenFileNames(self, "选择 Word 文档", "", "Word 文档 (*.docx)")
        if files: self.word_edit.setText(";".join(files))
    def select_excel(self):
        f, _ = QtWidgets.QFileDialog.getOpenFileName(self, "选择 Excel 文件", "", "Excel 文件 (*.xlsx)")
        if f: self.excel_edit.setText(f)
    def logmsg(self, s): self.log.appendPlainText(s); print(s)
    def preview_first(self):
        files = [p for p in self.word_edit.text().split(";") if p.strip()]
        if not files:
            self.logmsg("请先选择 Word 文件用于预览。"); return
        first = files[0]
        if not os.path.exists(first):
            self.logmsg("第一个文件不存在"); return
        try:
            data = extract_fields_from_docx(first)
            self.logmsg(f"预览 — {os.path.basename(first)}:")
            for k, v in data.items():
                self.logmsg(f"  {k} -> {v}")
        except Exception as e:
            self.logmsg(f"预览失败: {e}")
    def start_processing(self):
        files = [p for p in self.word_edit.text().split(";") if p.strip()]
        excel = self.excel_edit.text().strip()
        sheet = self.sheet_edit.text().strip() or None
        if not files: self.logmsg("请选择 Word 文件"); return
        if not excel: self.logmsg("请选择 Excel 文件"); return
        if not os.path.exists(excel): self.logmsg("Excel 文件不存在"); return
        self.start_btn.setEnabled(False); self.stop_btn.setEnabled(True); self.preview_btn.setEnabled(False)
        self.worker = Worker(files, excel, sheet)
        self.worker.log.connect(self.logmsg)
        self.worker.progress.connect(lambda a,b: self.progress.setValue(int(a*100/b) if b else 0))
        self.worker.finished.connect(self.on_finished)
        self.worker.start()
    def stop_processing(self):
        if self.worker:
            self.worker.stop(); self.logmsg("已请求停止...")
    def on_finished(self, processed, total):
        self.logmsg(f"完成：{processed}/{total}")
        self.start_btn.setEnabled(True); self.stop_btn.setEnabled(False); self.preview_btn.setEnabled(True)
        self.progress.setValue(100 if total and processed==total else int(processed*100/total) if total else 0)
        self.worker = None

# -------------------------
# entrypoint
# -------------------------
def main():
    app = QtWidgets.QApplication(sys.argv)
    w = MainWindow(); w.show()
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()
